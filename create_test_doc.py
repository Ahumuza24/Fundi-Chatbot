#!/usr/bin/env python3
from docx import Document

def create_test_docx():
    print("Creating test DOCX file...")
    
    # Create a new Document
    doc = Document()
    
    # Add a title
    doc.add_heading('Test Document for RAG Chatbot', 0)
    
    # Add some paragraphs
    doc.add_paragraph('This is a test document for the RAG chatbot system.')
    doc.add_paragraph('It contains sample text that can be used to test the document upload and processing functionality.')
    
    # Add a section with more content
    doc.add_heading('Sample Content', level=1)
    doc.add_paragraph('This document contains information about artificial intelligence and machine learning.')
    doc.add_paragraph('Machine learning is a subset of artificial intelligence that focuses on algorithms and statistical models.')
    doc.add_paragraph('These algorithms enable computers to perform tasks without being explicitly programmed.')
    
    # Add a list
    doc.add_heading('Key Features', level=2)
    features = [
        'Natural Language Processing',
        'Computer Vision',
        'Speech Recognition',
        'Recommendation Systems',
        'Autonomous Vehicles'
    ]
    
    for feature in features:
        doc.add_paragraph(feature, style='List Bullet')
    
    # Save the document
    doc.save('test_document.docx')
    print("✓ Test DOCX file created: test_document.docx")

if __name__ == "__main__":
    create_test_docx() 